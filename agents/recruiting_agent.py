import os
import re
import PyPDF2
import warnings

# Suppress warnings about python-docx
warnings.filterwarnings("ignore", category=UserWarning)

try:
    import python_docx as docx
except ImportError:
    try:
        import docx
    except ImportError:
        docx = None

import pandas as pd
from typing import Dict, Any, List
from models.embedding_model import EmbeddingModel
from utils.database_manager import DatabaseManager
from utils.logger import JobScreeningLogger

class RecruitingAgent:
    def __init__(self, embedding_model: EmbeddingModel, db_manager: DatabaseManager):
        """
        Initialize Recruiting Agent
        
        :param embedding_model: Embedding model for text processing
        :param db_manager: Database manager for storing candidate info
        """
        self.embedding_model = embedding_model
        self.db = db_manager
        self.logger = JobScreeningLogger()
        
        # Load dataset
        self.dataset_path = os.path.join(os.path.dirname(__file__), 'dataset.csv')
        try:
            # Try reading the file with different encodings and delimiters
            try:
                self.dataset = pd.read_csv(self.dataset_path, encoding='utf-8')
            except:
                try:
                    self.dataset = pd.read_csv(self.dataset_path, encoding='latin-1')
                except:
                    self.dataset = pd.read_csv(self.dataset_path, encoding='ISO-8859-1')
            
            # If no columns, treat the entire file as a single job description
            if len(self.dataset.columns) <= 1:
                with open(self.dataset_path, 'r', encoding='utf-8') as f:
                    job_description = f.read()
                self.dataset = pd.DataFrame({
                    'job_description': [job_description]
                })
        
        except Exception as e:
            self.logger.log_error('RecruitingAgent.init', f"Failed to load dataset: {e}")
            self.dataset = pd.DataFrame()

    def extract_text_from_resume(self, resume_path: str) -> str:
        """
        Extract text from resume file
        
        :param resume_path: Path to the resume file
        :return: Extracted text from the resume
        """
        try:
            # Extract text based on file extension
            file_extension = os.path.splitext(resume_path)[1].lower()
            
            if file_extension == '.pdf':
                # PDF text extraction
                with open(resume_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ' '.join(page.extract_text() for page in pdf_reader.pages)
            
            elif file_extension in ['.docx', '.doc']:
                # Check if docx is available
                if docx is None:
                    self.logger.log_error('RecruitingAgent.extract_text_from_resume', 'python-docx library not installed')
                    return ''
                
                # Word document text extraction
                doc = docx.Document(resume_path)
                text = ''
                for para in doc.paragraphs:
                    text += para.text + '\n'
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + '\n'
            
            else:
                # Plain text or unsupported format
                with open(resume_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            
            return text
        
        except Exception as e:
            self.logger.log_error('RecruitingAgent.extract_text_from_resume', e)
            return ''

    def _extract_experience(self, resume_text: str) -> List[Dict[str, str]]:
        """
        Extract work experience from resume text
        
        :param resume_text: Resume text
        :return: List of work experiences
        """
        try:
            # Use regex to find potential work experiences
            experience_patterns = [
                r'(\w+)\s*(?:at|@)\s*(\w+(?:\s+\w+)*)\s*(?:from|for)?\s*(\d{4}(?:\s*-\s*\d{4})?)',
                r'(\w+(?:\s+\w+)*)\s*(?:experience|worked)\s*(?:at|@)\s*(\w+(?:\s+\w+)*)',
            ]
            
            experiences = []
            for pattern in experience_patterns:
                matches = re.findall(pattern, resume_text, re.IGNORECASE)
                for match in matches:
                    experiences.append({
                        'role': match[0],
                        'company': match[1],
                        'duration': match[2] if len(match) > 2 else 'Unknown'
                    })
            
            return experiences
        
        except Exception as e:
            self.logger.log_error('RecruitingAgent._extract_experience', e)
            return []

    def _extract_education(self, resume_text: str) -> List[Dict[str, str]]:
        """
        Extract education from resume text
        
        :param resume_text: Resume text
        :return: List of educational qualifications
        """
        try:
            # Use regex to find potential educational qualifications
            education_patterns = [
                r'(\w+(?:\s+\w+)*)\s*(?:degree|graduated)\s*(?:from)?\s*(\w+(?:\s+\w+)*)\s*(?:in)?\s*(\w+(?:\s+\w+)*)',
                r'(\w+(?:\s+\w+)*)\s*(?:at|@)\s*(\w+(?:\s+\w+)*)\s*(?:graduated|completed)',
            ]
            
            educations = []
            for pattern in education_patterns:
                matches = re.findall(pattern, resume_text, re.IGNORECASE)
                for match in matches:
                    educations.append({
                        'degree': match[0],
                        'institution': match[1],
                        'year': match[2] if len(match) > 2 else 'Unknown'
                    })
            
            return educations
        
        except Exception as e:
            self.logger.log_error('RecruitingAgent._extract_education', e)
            return []

    def process_candidate_resume(self, resume_path: str, candidate_name: str, email: str) -> int:
        """
        Process candidate resume and store information
        
        :param resume_path: Path to resume file
        :param candidate_name: Name of the candidate
        :param email: Candidate's email
        :return: Candidate ID
        """
        # Extract resume text
        resume_text = self.extract_text_from_resume(resume_path)
        
        # Extract experiences and education
        experiences = self._extract_experience(resume_text)
        education = self._extract_education(resume_text)
        
        # Log extraction details
        self.logger.log_candidate_extraction(
            candidate_id=hash(email),  # Use email hash as temporary ID
            skills=[exp.get('role', '') for exp in experiences],
            experience=[exp.get('company', '') for exp in experiences]
        )
        
        # Store candidate information in database
        candidate_id = self.db.store_candidate({
            'name': candidate_name,
            'email': email,
            'resume_text': resume_text,
            'experiences': experiences,
            'education': education
        })
        
        return candidate_id